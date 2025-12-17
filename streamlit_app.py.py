import streamlit as st
from pdf2docx import Converter
import os
import tempfile
import zipfile
from pathlib import Path
import fitz  # PyMuPDF for text detection and image rendering
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Page config
st.set_page_config(page_title="Arabic PDF → Word Converter", page_icon="📄")
st.title("🇸🇦 Arabic PDF to Word Converter")
st.markdown("""
Upload Arabic PDFs → get editable Word files.  
Supports layout preservation, tables, images, and RTL text.
""")

# Sidebar options
st.sidebar.header("Conversion Options")
convert_all_pages = st.sidebar.checkbox("Convert all pages", value=True)

if not convert_all_pages:
    col1, col2 = st.sidebar.columns(2)
    start_page = col1.number_input("Start page", min_value=1, value=1, step=1)
    end_page = col2.number_input("End page", min_value=1, value=10, step=1)
    if start_page > end_page:
        st.sidebar.error("Start page must be ≤ End page")
        st.stop()
else:
    start_page = None
    end_page = None

image_fallback = st.sidebar.checkbox(
    "Preserve exact layout (embed pages as images)",
    value=False,
    help="Best for scanned or complex PDFs – 100% visual match, but text not editable/searchable"
)

# File uploader
uploaded_files = st.file_uploader(
    "Choose PDF file(s)",
    type="pdf",
    accept_multiple_files=True,
    help="Select multiple PDFs for batch conversion"
)

if not uploaded_files:
    st.info("👆 Upload one or more PDFs to begin.")
    st.stop()

# Temporary directory
with tempfile.TemporaryDirectory() as temp_dir:
    temp_path = Path(temp_dir)
    output_files = []

    total_tasks = len(uploaded_files)
    progress_bar = st.progress(0)
    status_text = st.empty()

    for idx, uploaded_file in enumerate(uploaded_files):
        filename_base = Path(uploaded_file.name).stem
        input_pdf_path = temp_path / uploaded_file.name
        with open(input_pdf_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # PDF Preview (first file only, or in expander for multiples)
        if idx == 0 or len(uploaded_files) == 1:
            with st.expander("📄 Preview uploaded PDF", expanded=True):
                st.pdf(uploaded_file.getvalue(), height=800)

        # Scanned detection
        doc = fitz.open(stream=uploaded_file.getvalue(), filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        is_scanned = len(text.strip()) < 100  # Heuristic: very little text = likely scanned
        doc.close()

        if is_scanned:
            st.warning(
                f"⚠️ **{uploaded_file.name}** appears to be a scanned/image-based PDF. "
                "Text may not be editable. Enable 'Preserve exact layout' for perfect visual results."
            )

        output_docx_path = temp_path / f"{filename_base}.docx"

        status_text.text(f"Processing {idx+1}/{total_tasks}: {uploaded_file.name}")

        try:
            if image_fallback:
                # Image fallback mode
                word_doc = Document()
                word_doc.add_heading(f"Converted from: {uploaded_file.name}", level=1)

                total_pages = len(doc) if 'doc' in locals() else fitz.open(str(input_pdf_path)).page_count
                page_progress = st.empty()

                for page_num in range(total_pages):
                    if not convert_all_pages and not (start_page <= page_num+1 <= end_page):
                        continue

                    page_progress.progress((page_num + 1) / total_pages)
                    page_progress.text(f"Rendering page {page_num+1} of {total_pages} as image")

                    page = doc.load_page(page_num)
                    mat = fitz.Matrix(300/72, 300/72)  # High DPI
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img_bytes = pix.tobytes("png")
                    img_io = BytesIO(img_bytes)

                    word_doc.add_paragraph().add_run().add_picture(img_io, width=Inches(6.5))
                    word_doc.add_page_break()

                word_doc.save(str(output_docx_path))
                page_progress.empty()
            else:
                # Normal pdf2docx conversion
                cv = Converter(str(input_pdf_path))
                total_pages = len(cv.pages) if hasattr(cv, 'pages') else doc.page_count

                def page_callback(current_page: int):
                    progress_percent = current_page / total_pages
                    st.progress(progress_percent)
                    st.caption(f"Converting page {current_page} of {total_pages}")

                if convert_all_pages:
                    cv.convert(str(output_docx_path), callback=page_callback)
                else:
                    cv.convert(str(output_docx_path), start=start_page-1, end=end_page, callback=page_callback)
                cv.close()

            output_files.append((output_docx_path, f"{filename_base}.docx"))

        except Exception as e:
            st.error(f"Failed to convert **{uploaded_file.name}**: {str(e)}")
            continue

        progress_bar.progress((idx + 1) / total_tasks)

    status_text.text("All done!")
    progress_bar.empty()

    # Downloads
    if len(output_files) == 1:
        docx_path, display_name = output_files[0]
        with open(docx_path, "rb") as f:
            st.download_button(
                "📥 Download Word file",
                f,
                file_name=display_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        zip_path = temp_path / "converted_files.zip"
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for docx_path, arcname in output_files:
                zipf.write(docx_path, arcname)
        with open(zip_path, "rb") as f:
            st.download_button(
                f"📦 Download all {len(output_files)} Word files (ZIP)",
                f,
                file_name="arabic_converted_files.zip",
                mime="application/zip"
            )

    st.success("Conversion complete!")
    st.balloons()

    # Feedback
    st.markdown("---")
    st.caption("Was this tool helpful?")
    feedback = st.feedback("thumbs")
    if feedback == 0:
        st.toast("😢 Sorry to hear that! Feedback helps us improve.")
    elif feedback == 1:
        st.toast("👍 Thank you! Glad it worked well.")

st.markdown("Made with ❤️ for Arabic document lovers")
