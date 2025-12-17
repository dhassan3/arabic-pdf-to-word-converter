import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from pathlib import Path
import tempfile
import zipfile
from io import BytesIO

st.set_page_config(page_title="Arabic PDF → Word Converter", page_icon="📄")
st.title("🇸🇦 Arabic PDF to Word Converter")
st.markdown("Upload Arabic PDFs → get editable Word files. Layout & RTL preserved as best as possible.")

# Sidebar options
st.sidebar.header("Options")
convert_all_pages = st.sidebar.checkbox("Convert all pages", value=True)
if not convert_all_pages:
    col1, col2 = st.sidebar.columns(2)
    start_page = col1.number_input("Start page", min_value=1, value=1, step=1)
    end_page = col2.number_input("End page", min_value=1, value=10, step=1)
    if start_page > end_page:
        st.sidebar.error("Start page must be ≤ End page")
        st.stop()

image_fallback = st.sidebar.checkbox(
    "Exact layout mode (embed pages as images)",
    value=False,
    help="Recommended for scanned PDFs – perfect visual match (text not editable)"
)

# File uploader
uploaded_files = st.file_uploader("Choose PDF file(s)", type="pdf", accept_multiple_files=True)

if not uploaded_files:
    st.info("👆 Upload one or more PDFs to start.")
    st.stop()

# Processing
with tempfile.TemporaryDirectory() as temp_dir:
    temp_path = Path(temp_dir)
    output_files = []
    overall_progress = st.progress(0)
    status_text = st.empty()

    for idx, uploaded_file in enumerate(uploaded_files):
        filename_base = Path(uploaded_file.name).stem
        input_pdf_path = temp_path / uploaded_file.name
        with open(input_pdf_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Simple preview (first file or single)
        if idx == 0 or len(uploaded_files) == 1:
            with st.expander("👁️ View uploaded PDF (download to open)"):
                st.download_button(
                    "Download PDF for local viewing",
                    uploaded_file.getvalue(),
                    file_name=uploaded_file.name,
                    mime="application/pdf"
                )

        # Scanned detection
        doc = fitz.open(stream=uploaded_file.getvalue(), filetype="pdf")
        text_chars = sum(len(page.get_text()) for page in doc)
        if text_chars < 100:
            st.warning(f"⚠️ **{uploaded_file.name}** appears scanned/image-based. Enable 'Exact layout mode' for best results.")

        output_docx_path = temp_path / f"{filename_base}.docx"
        status_text.text(f"Processing {idx+1}/{len(uploaded_files)}: {uploaded_file.name}")

        try:
            word_doc = Document()
            word_doc.add_heading(f"Converted from: {uploaded_file.name}", level=1)

            total_pages = doc.page_count
            page_progress = st.empty()

            for page_num in range(total_pages):
                if not convert_all_pages and not (start_page <= page_num + 1 <= end_page):
                    continue

                page_progress.progress((page_num + 1) / total_pages)
                page_progress.text(f"Page {page_num + 1} of {total_pages}")

                page = doc.load_page(page_num)

                if image_fallback:
                    # Embed as high-quality image
                    mat = fitz.Matrix(300/72, 300/72)  # 300 DPI
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img_bytes = pix.tobytes("png")
                    word_doc.add_paragraph().add_run().add_picture(BytesIO(img_bytes), width=Inches(6.5))
                    word_doc.add_page_break()
                else:
                    # Basic text extraction with RTL hint
                    blocks = page.get_text("dict")["blocks"]
                    for block in blocks:
                        if "lines" in block:
                            para_text = ""
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    para_text += span["text"]
                            para_text = para_text.strip()
                            if para_text:
                                p = word_doc.add_paragraph(para_text)
                                if any("\u0600" <= c <= "\u06FF" for c in para_text):  # Arabic chars
                                    p.paragraph_format.right_to_left = True

            word_doc.save(str(output_docx_path))
            page_progress.empty()
            output_files.append((output_docx_path, f"{filename_base}.docx"))

        except Exception as e:
            st.error(f"Failed on {uploaded_file.name}: {str(e)}")
            continue

        overall_progress.progress((idx + 1) / len(uploaded_files))

    status_text.text("All complete!")
    overall_progress.empty()

    # Download section
    if len(output_files) == 1:
        with open(output_files[0][0], "rb") as f:
            st.download_button(
                "📥 Download Word file",
                f,
                file_name=output_files[0][1],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        zip_path = temp_path / "converted_files.zip"
        with zipfile.ZipFile(zip_path, "w") as z:
            for path, name in output_files:
                z.write(path, name)
        with open(zip_path, "rb") as f:
            st.download_button(
                "📦 Download all Word files (ZIP)",
                f,
                file_name="converted_arabic_files.zip",
                mime="application/zip"
            )

    st.success("Success! Files ready.")
    st.balloons()

    # Feedback
    st.caption("Was this helpful?")
    st.feedback("thumbs")

st.markdown("Made with ❤️ for Arabic documents")
